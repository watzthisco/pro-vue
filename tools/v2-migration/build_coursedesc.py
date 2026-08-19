#!/usr/bin/env python3
"""Refresh the Version 2.0 course description for the Vue 3 / 2026 course."""
import docx

PATH = '/home/user/pro-vue/presentation/Version 2.0/Vue-CourseDescription.docx'

REPLACEMENTS = {
    'Version 2.0, September 2025': 'Version 2.0, 2026',
    'Use Vuex for maintaining state in a Vue.js application':
        'Use Pinia for maintaining state in a Vue.js application',
    'State management with Vuex': 'State management with Pinia',
    'Using Vuex': 'Using Pinia',
    'Lab 14: Implementing Vuex': 'Lab 14: Implementing Pinia',
    'Lab 16: AJAX with Vuex': 'Lab 16: AJAX with Pinia',
    'Using Filters': 'Binding HTML Classes',
    'The Vue Instance': 'Creating and Mounting an App',
    'Instance Properties and Methods': 'The Composition API and <script setup>',
    'Vue Instance Lifecycle': 'Component Lifecycle',
    'Gain a deep knowledge of Vue.js components':
        'Gain a deep knowledge of Vue.js components and the Composition API',
    'Learn to create test suites for Vue':
        'Learn to create test suites for Vue with Vitest',
    'Vue State': 'Reactive State with ref and reactive',
    'Computed Properties': 'Computed Properties',
    'Creating an SPA with Vue-Router': 'Creating an SPA with Vue Router',
}

# Paragraphs inserted after a given outline entry, to cover material the Vue 3
# course adds.
ADDITIONS = {
    'Using Lifecycle Hooks': ['Composables (replacing Mixins)'],
    'Vue.js Directives': ['Reactivity Fundamentals'],
}

doc = docx.Document(PATH)

changed = 0
for p in doc.paragraphs:
    text = p.text.strip()
    if text in REPLACEMENTS and REPLACEMENTS[text] != text:
        for r in list(p.runs):
            r._r.getparent().remove(r._r)
        p.add_run(REPLACEMENTS[text])
        changed += 1

# Insert the additions after their anchor paragraphs.
import copy
added = 0
for anchor_text, new_texts in ADDITIONS.items():
    for p in doc.paragraphs:
        if p.text.strip() == anchor_text:
            el = p._p
            for t in reversed(new_texts):
                new_el = copy.deepcopy(el)
                for child in list(new_el):
                    if child.tag.endswith('}r'):
                        new_el.remove(child)
                el.addnext(new_el)
                from docx.text.paragraph import Paragraph
                Paragraph(new_el, doc._body).add_run(t)
                added += 1
            break

doc.save(PATH)
print(f'updated {changed} entries, added {added}')
