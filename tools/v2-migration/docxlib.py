"""Small helpers for surgically rewriting paragraph ranges in a python-docx document."""
import copy


def clone_paragraph_shell(model_para):
    """A copy of `model_para` with its runs stripped, so styling/numbering survives."""
    new_p = copy.deepcopy(model_para._p)
    for child in list(new_p):
        if child.tag.endswith('}r') or child.tag.endswith('}hyperlink'):
            new_p.remove(child)
    return new_p


def find_style_model(doc, style_name):
    """The first paragraph using `style_name`, used as a template for new paragraphs."""
    for p in doc.paragraphs:
        if p.style.name == style_name:
            return p
    raise KeyError(f'no paragraph uses style {style_name!r}')


def replace_range(doc, start, end, new_items, models):
    """Replace paragraphs [start, end] with `new_items`, a list of (style, text).

    Returns the number of paragraphs written. Ranges must be applied from the
    end of the document backwards so earlier indices stay valid.
    """
    paras = doc.paragraphs
    anchor = paras[start]._p
    parent = anchor.getparent()

    written = []
    for style, text in new_items:
        model = models[style]
        p_el = clone_paragraph_shell(model)
        parent.insert(parent.index(anchor), p_el)
        written.append(p_el)

    # Remove the originals now that the replacements are in place.
    for p in paras[start:end + 1]:
        p._p.getparent().remove(p._p)

    # Fill in the text (paragraph objects must be rebuilt after the DOM edits).
    from docx.text.paragraph import Paragraph
    for p_el, (style, text) in zip(written, new_items):
        para = Paragraph(p_el, doc._body)
        para.style = doc.styles[style]
        if text:
            para.add_run(text)
    return len(new_items)


def set_text(doc, index, text, style=None):
    """Replace the text of a single paragraph, keeping (or setting) its style."""
    para = doc.paragraphs[index]
    for r in list(para.runs):
        r._r.getparent().remove(r._r)
    if style:
        para.style = doc.styles[style]
    if text:
        para.add_run(text)
