#!/usr/bin/env python3
"""Build Professional-Vue-v2.0.docx from the v1.5 lab manual."""
import os, sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

import docx
from docxlib import replace_range, find_style_model, set_text
import labcontent as C

ROOT = '/home/user/pro-vue'
SRC = os.path.join(ROOT, 'presentation/Version 1.5/Professional-Vue-v1.5.docx')
DST = os.path.join(ROOT, 'presentation/Version 2.0/Professional-Vue-v2.0.docx')

doc = docx.Document(SRC)

STYLES = ['Normal', 'Heading 1', 'Heading 2', 'LabNL', 'LabNL Sub', 'LabNL Code',
          'Note', 'Code', 'LabBullet', 'List Paragraph']
models = {s: find_style_model(doc, s) for s in STYLES}

# (start, end, new_items) -- inclusive paragraph indices in the v1.5 document.
EDITS = [
    (9,    11,   C.FRONT),
    (77,   107,  C.SETUP),
    (118,  144,  C.LAB01),
    (145,  169,  C.LAB02),
    (177,  290,  C.LAB03),
    (357,  388,  C.LAB05),
    (494,  499,  C.LAB08_CODE),
    (506,  583,  C.LAB09),
    (584,  618,  C.LAB10),
    (623,  716,  C.LAB11_SCRIPTS),
    (767,  792,  C.LAB13),
    (793,  802,  C.LAB14),
    (803,  828,  C.LAB15),
    (829,  899,  C.LAB16),
    (900,  935,  C.LAB17),
    (940,  945,  C.LAB19),
    (946,  1109, C.LAB20),
    (1110, 1220, C.LAB21),
    (1221, 1259, C.LAB22),
    (1260, 1265, C.LAB23),
    (1266, 1282, C.LAB24),
]

# Sanity-check that the ranges still line up with the source document before
# any of them are applied.
before = [(i, p.style.name, p.text[:60]) for i, p in enumerate(doc.paragraphs)]
for start, end, items in EDITS:
    assert 0 <= start <= end < len(before), f'range {start}-{end} out of bounds'

# Apply back to front so earlier indices stay valid.
for start, end, items in sorted(EDITS, reverse=True):
    n = replace_range(doc, start, end, items, models)
    print(f'  {start:5d}-{end:<5d} -> {n:3d} paragraphs')

# Single-paragraph rewrites in sections that are otherwise unchanged.
tweaked = 0
for para in doc.paragraphs:
    text = para.text.strip()
    if not text:
        continue
    for prefix, replacement in C.TWEAKS:
        if text.startswith(prefix):
            for r in list(para.runs):
                r._r.getparent().remove(r._r)
            para.add_run(replacement)
            tweaked += 1
            break
print(f'  applied {tweaked} single-paragraph tweaks')

# Rename the table-of-contents entries whose lab titles changed. The page
# numbers go stale; Word refreshes them when the field is updated (F9).
renamed = 0
for p in doc.paragraphs:
    if not p.style.name.startswith('toc'):
        continue
    for run in p.runs:
        if run.text in C.TOC_RENAMES:
            run.text = C.TOC_RENAMES[run.text]
            renamed += 1
print(f'  renamed {renamed} TOC entries')

os.makedirs(os.path.dirname(DST), exist_ok=True)
doc.save(DST)
print('wrote', DST)
